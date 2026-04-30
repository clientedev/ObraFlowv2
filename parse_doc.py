from docx import Document
from docx.shared import RGBColor

doc = Document('Aplicativo ELP - Validação - 29.04.docx')

results = []
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    
    # Check for red text in runs
    has_red = False
    red_texts = []
    for run in para.runs:
        if run.font.color and run.font.color.rgb:
            r_val = run.font.color.rgb[0]
            g_val = run.font.color.rgb[1]
            b_val = run.font.color.rgb[2]
            if r_val > 180 and g_val < 100 and b_val < 100:
                has_red = True
                red_texts.append(run.text)
    
    style = para.style.name if para.style else 'None'
    marker = '[RED]' if has_red else '     '
    print(marker + ' [' + style + '] ' + text)

# Also check tables
print("\n\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- Table {i+1} ---")
    for row in table.rows:
        row_texts = []
        row_has_red = False
        for cell in row.cells:
            cell_text = cell.text.strip()
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.color and run.font.color.rgb:
                        r_val = run.font.color.rgb[0]
                        g_val = run.font.color.rgb[1]
                        b_val = run.font.color.rgb[2]
                        if r_val > 180 and g_val < 100 and b_val < 100:
                            row_has_red = True
            row_texts.append(cell_text)
        marker = '[RED]' if row_has_red else '     '
        print(marker + ' | '.join(row_texts))
